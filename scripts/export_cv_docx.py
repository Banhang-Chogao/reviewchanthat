#!/usr/bin/env python3
"""
Export CV to professional DOCX format using python-docx.
Requires: python-docx

Usage:
  python3 export_cv_docx.py --data '{"fullName":"Nguyen Van A","email":"..."}' --output cv.docx

Features:
  - A4 page setup, professional recruiter layout
  - Unicode support
  - Proper styles, ATS-friendly
"""

import argparse
import json
import os
import sys

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


PRIMARY_HEX = '00A7A0'
DARK_HEX = '222222'
GRAY_HEX = '666666'

def hex_to_rgb(h):
    h = h.lstrip('#')
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


class CVDocx:
    def __init__(self, data, output_path):
        self.data = data
        self.output_path = output_path
        self.doc = Document()
        self._setup_page()
        self._setup_styles()

    def _setup_page(self):
        section = self.doc.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    def _setup_styles(self):
        style = self.doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(10)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.15

    def _add_heading_custom(self, text, size=16, bold=True, color=DARK_HEX, space_before=0, space_after=4):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = hex_to_rgb(color)
        return p

    def _add_section_title(self, text):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text.upper())
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = hex_to_rgb(PRIMARY_HEX)
        # add bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="{PRIMARY_HEX}"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)
        return p

    def _add_entry_header(self, title, subtitle='', date=''):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        if title:
            run = p.add_run(title)
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.color.rgb = hex_to_rgb(DARK_HEX)
        if date:
            run = p.add_run(f'    {date}')
            run.font.size = Pt(8)
            run.font.color.rgb = hex_to_rgb(GRAY_HEX)
        if subtitle:
            p2 = self.doc.add_paragraph()
            p2.paragraph_format.space_after = Pt(1)
            run = p2.add_run(subtitle)
            run.font.size = Pt(9)
            run.font.bold = True
            run.font.color.rgb = hex_to_rgb(PRIMARY_HEX)

    def _add_body(self, text, size=9, color=DARK_HEX):
        if not text: return
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.color.rgb = hex_to_rgb(color)

    def _add_bullet_list(self, items, size=9):
        for item in items:
            if not item.strip(): continue
            p = self.doc.add_paragraph(style='List Bullet')
            p.clear()
            run = p.add_run(item.strip())
            run.font.size = Pt(size)
            run.font.color.rgb = hex_to_rgb(DARK_HEX)

    def _add_skills(self, skills):
        if not skills: return
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        for i, s in enumerate(skills):
            run = p.add_run(f'  {s}  ')
            run.font.size = Pt(8)
            run.font.bold = True
            run.font.color.rgb = hex_to_rgb('FFFFFF')
            run.font.highlight_color = None
            # Add shading to run
            rPr = run._r.get_or_add_rPr()
            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{PRIMARY_HEX}" w:val="clear"/>')
            rPr.append(shd)

    def build(self):
        d = self.data

        # Name
        self._add_heading_custom(d.get('fullName', 'Your Name'), size=18)

        # Title
        title_text = d.get('objective', '')
        if title_text:
            title_text = title_text.split('.')[0]
            self._add_heading_custom(title_text, size=11, color=PRIMARY_HEX)

        # Contact
        contact_parts = []
        if d.get('email'): contact_parts.append(d['email'])
        if d.get('phone'): contact_parts.append(d['phone'])
        if d.get('address'): contact_parts.append(d['address'])
        if contact_parts:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(' | '.join(contact_parts))
            run.font.size = Pt(8)
            run.font.color.rgb = hex_to_rgb(GRAY_HEX)

        # Horizontal line
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(6)
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:bottom w:val="single" w:sz="12" w:space="1" w:color="{PRIMARY_HEX}"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)

        # Objective
        if d.get('objective'):
            self._add_section_title('Career Objective')
            self._add_body(d['objective'])

        # Education
        if d.get('school') or d.get('degree'):
            self._add_section_title('Education')
            if d.get('school'):
                self._add_entry_header(d['school'], '', f"{d.get('eduStart', '')} - {d.get('eduEnd', '')}")
            if d.get('degree'):
                self._add_body(d['degree'], size=9, color=PRIMARY_HEX)
            if d.get('eduGpa'):
                self._add_body(f"GPA: {d['eduGpa']}")

        # Experience
        exps = d.get('experiences', [])
        valid_exps = [e for e in exps if e.get('company') or e.get('position')]
        if valid_exps:
            self._add_section_title('Work Experience')
            for e in valid_exps:
                date_str = f"{e.get('start', '')} - {e.get('end', '')}"
                self._add_entry_header(e.get('position', ''), e.get('company', ''), date_str)
                if e.get('desc'):
                    self._add_body(e['desc'])

        # Projects
        projects = d.get('projects', [])
        valid_proj = [p for p in projects if p.get('name') or p.get('desc')]
        if valid_proj:
            self._add_section_title('Projects')
            for p in valid_proj:
                if p.get('name'):
                    self._add_heading_custom(p['name'], size=10, color=DARK_HEX, space_before=2, space_after=1)
                if p.get('desc'):
                    self._add_body(p['desc'])
                tech_parts = []
                if p.get('tech'): tech_parts.append(p['tech'])
                if p.get('url'): tech_parts.append(p['url'])
                if tech_parts:
                    p_text = ' | '.join(tech_parts)
                    self._add_body(p_text, size=8, color=PRIMARY_HEX)

        # Skills
        skills = d.get('skills', [])
        if skills:
            self._add_section_title('Skills')
            self._add_skills(skills)

        # Certifications
        certs = d.get('certifications', '')
        if certs:
            self._add_section_title('Certifications')
            self._add_bullet_list([c.strip() for c in certs.split('\n') if c.strip()])

        # Languages
        langs = d.get('languages', '')
        if langs:
            self._add_section_title('Languages')
            self._add_bullet_list([l.strip() for l in langs.split('\n') if l.strip()])

        # Awards
        awards = d.get('awards', '')
        if awards:
            self._add_section_title('Awards')
            self._add_bullet_list([a.strip() for a in awards.split('\n') if a.strip()])

        # Activities
        activities = d.get('activities', '')
        if activities:
            self._add_section_title('Activities')
            self._add_bullet_list([a.strip() for a in activities.split('\n') if a.strip()])

        # References
        refs = d.get('references', '')
        if refs:
            self._add_section_title('References')
            self._add_body(refs)

        # Additional
        add = d.get('additional', '')
        if add:
            self._add_section_title('Additional Information')
            self._add_body(add)

        self.doc.save(self.output_path)
        print(f'DOCX generated: {self.output_path}')


def main():
    parser = argparse.ArgumentParser(description='Export CV to DOCX')
    parser.add_argument('--data', required=True, help='JSON string with CV data')
    parser.add_argument('--output', default='cv.docx', help='Output DOCX path')
    args = parser.parse_args()

    try:
        data = json.loads(args.data)
    except json.JSONDecodeError as e:
        print(f'Invalid JSON: {e}')
        sys.exit(1)

    cv = CVDocx(data, args.output)
    cv.build()


if __name__ == '__main__':
    main()
